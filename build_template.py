"""
Builds the master fee-proposal template (.docx) for the Concept Engineers
proposal generation system. The output is a docxtpl/Jinja2 template: it
contains placeholder variables and conditional logic that the Python
generation engine fills at runtime.

Key docxtpl mechanics used:
  {{ var }}                  -> inline variable substitution
  {%p if cond %}...{%p endif %}  -> removes the WHOLE paragraph carrying the
                                    tag, so hidden sections leave NO blank lines
  {%tr for x in list %}...{%tr endfor %} -> repeats / removes a whole TABLE ROW
This is what satisfies the "no blank pages or layout issues" requirement.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x33, 0x55)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x59, 0x59, 0x59)
LIGHT = "EAF1F8"
RULE = "2E75B6"

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, bold, before, after in [
    ("H1", 15, NAVY, True, 14, 6),
    ("H2", 12, ACCENT, True, 10, 4),
]:
    st = doc.styles.add_style(name, 1)  # paragraph style
    st.base_style = doc.styles["Normal"]
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = color
    st.font.name = "Calibri"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# ---------- page setup ----------
sec = doc.sections[0]
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(1.8)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def bottom_rule(paragraph, color=RULE, size=8):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def para(text="", style=None, size=None, color=None, bold=False, italic=False,
         align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        r.bold = bold
        r.italic = italic
    return p


def tag(text):
    """A paragraph that holds only a docxtpl control tag (gets removed)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.add_run(text)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def add_hyperlink(paragraph, url, text, color="2E75B6"):
    """Insert a real external hyperlink as a sibling of the runs (valid XML)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# =====================================================================
# HEADER (firm letterhead) and FOOTER (page numbers)
# =====================================================================
hdr = sec.header
hp = hdr.paragraphs[0]
hr = hp.add_run("{{ firm_name }}")
hr.font.size = Pt(16)
hr.font.bold = True
hr.font.color.rgb = NAVY
hp2 = hdr.add_paragraph()
hr2 = hp2.add_run("Consulting Civil & Structural Engineers    |    ABN {{ firm_abn }}    |    {{ firm_web }}")
hr2.font.size = Pt(8)
hr2.font.color.rgb = GREY
bottom_rule(hp2)

ftr = sec.footer
fp = ftr.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("{{ document_type }}  |  Ref {{ proposal_ref }}  |  {{ revision }}  |  Page ")
run.font.size = Pt(8)
run.font.color.rgb = GREY
# PAGE field
fld1 = OxmlElement("w:fldSimple")
fld1.set(qn("w:instr"), "PAGE")
fp._p.append(fld1)
run2 = fp.add_run(" of ")
run2.font.size = Pt(8)
run2.font.color.rgb = GREY
fld2 = OxmlElement("w:fldSimple")
fld2.set(qn("w:instr"), "NUMPAGES")
fp._p.append(fld2)

# =====================================================================
# TITLE BLOCK
# =====================================================================
para("{{ document_type | upper }}", style="H1")
meta = para()
mr = meta.add_run("Ref {{ proposal_ref }}    |    {{ revision }}    |    {{ proposal_date }}")
mr.font.size = Pt(9)
mr.font.color.rgb = GREY

# recipient
para("")
para("{{ contact_name }}", bold=True, after=0)
para("{{ company_name }}", after=0)
para("{{ project_address }}", after=0)
para("{{ contact_email }}", after=8)

# subject line
subj = para()
subj.add_run("RE:  ").bold = True
sr = subj.add_run("{{ document_type }} for {{ project_address }}")
sr.bold = True
sr.font.color.rgb = NAVY

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
para("Dear {{ contact_name }},", after=6)
para(
    "Thank you for the opportunity to provide this {{ document_type | lower }} for "
    "{{ company_name }}. We are pleased to submit our proposed scope of services and "
    "associated fees for the project at {{ project_address }}. This document sets out "
    "the services to be provided, our relevant capability, and the fee basis for the "
    "engagement."
)

# =====================================================================
# 2. PROJECT UNDERSTANDING
# =====================================================================
para("Project Understanding", style="H2")
para("{{ project_understanding }}")

# =====================================================================
# 3. SCOPE OF SERVICES  (conditional sub-sections)
# =====================================================================
para("Scope of Services", style="H2")
para("Our proposed scope for this engagement comprises the following components:")

# --- Development Approval (DA) ---
tag("{%p if is_da %}")
para("Development Approval (DA)", bold=True, color=ACCENT, after=2)
bullet("Preparation of civil engineering plans and reports to support the development application.")
bullet("Stormwater management strategy and concept drainage design.")
bullet("Earthworks, access and servicing concept to satisfy authority requirements.")
bullet("Liaison with the relevant assessment authority through the DA process.")
tag("{%p endif %}")

# --- Detailed Design (DD) ---
tag("{%p if is_dd %}")
para("Detailed Design (DD)", bold=True, color=ACCENT, after=2)
bullet("Detailed civil engineering design and documentation for construction.")
bullet("Stormwater, drainage, pavement and earthworks design to approved standards.")
bullet("Preparation of a construction-issue drawing set and specifications.")
bullet("Engineering certification of the detailed design.")
tag("{%p endif %}")

# --- UU Endorsed Consultant scope ---
tag("{%p if is_uu %}")
para("Urban Utilities (UU) Endorsed Consultant", bold=True, color=ACCENT, after=2)
bullet("Water and sewer design undertaken as an Urban Utilities endorsed consultant.")
bullet("Preparation and lodgement of the relevant UU application and design package.")
bullet("Coordination with Urban Utilities through to design acceptance.")
tag("{%p endif %}")

# --- UW Certifier scope ---
tag("{%p if is_uw %}")
para("Certifier Scope ({{ uw_level }})", bold=True, color=ACCENT, after=2)
bullet("Certification services provided in the capacity of {{ uw_level }} certifier.")
bullet("Review and certification of works against the approved design and standards.")
bullet("Issue of the relevant compliance and certification documentation.")
tag("{%p endif %}")

# --- UU works-only simplified note ---
tag("{%p if uu_works_only %}")
note = para()
note.add_run("Note: ").bold = True
note.add_run(
    "This engagement is for Urban Utilities works only. A simplified fee structure "
    "applies as set out in the fee schedule below."
).italic = True
tag("{%p endif %}")

# =====================================================================
# 4. CAPABILITY & EXPERIENCE  (1 of 9, selected by capability_profile)
# =====================================================================
para("Capability & Experience", style="H2")

capability_blocks = [
    ("residential", "Residential Subdivision",
     "Our team has delivered civil engineering for residential subdivisions ranging "
     "from small infill lots to large master-planned communities, covering DA support, "
     "detailed design, stormwater and servicing through to construction certification."),
    ("commercial", "Commercial & Industrial",
     "We provide civil engineering for commercial and industrial developments, including "
     "site grading, pavement design, stormwater quality and quantity management, and "
     "coordination with service authorities for fast-tracked delivery."),
    ("infrastructure", "Civil Infrastructure",
     "Our experience spans roads, drainage networks and public infrastructure, delivered "
     "to relevant authority standards with a focus on constructability and whole-of-life "
     "value."),
    ("stormwater", "Stormwater & Drainage",
     "We specialise in stormwater management, from catchment modelling and water quality "
     "treatment to detailed drainage design and authority approvals."),
]

for key, title, body in capability_blocks:
    tag("{%p if capability_profile == '" + key + "' %}")
    para(title, bold=True, color=ACCENT, after=2)
    para(body)
    tag("{%p endif %}")

# generic fallback so a page is always present
tag("{%p if capability_profile not in known_profiles %}")
para("{{ capability_title }}", bold=True, color=ACCENT, after=2)
para("{{ capability_body }}")
tag("{%p endif %}")

# NOTE marker for the remaining profiles (extend to all 9 by copying a block above)
# Remaining profiles to add: townhouse, mixeduse, civil_works, land_development, civil_certification

# =====================================================================
# 5. FEE SCHEDULE
# =====================================================================
para("Fee Schedule", style="H2")
para("Our proposed professional fees for the scope described above are as follows:")

# --- fixed fee table with looped rows ---
fee_tbl = doc.add_table(rows=1, cols=2)
fee_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
fee_tbl.style = "Table Grid"
hdr_cells = fee_tbl.rows[0].cells
for c, txt, w in [(hdr_cells[0], "Service Component", Cm(12)), (hdr_cells[1], "Fee (AUD, ex GST)", Cm(4.5))]:
    set_cell_bg(c, "1F3355")
    cp = c.paragraphs[0]
    rr = cp.add_run(txt)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rr.font.size = Pt(10)

# control row -> {% for %}  (whole row is removed by docxtpl)
fee_tbl.add_row().cells[0].paragraphs[0].add_run("{%tr for item in fee_items %}")
# body row (repeats once per fee item)
row = fee_tbl.add_row().cells
row[0].paragraphs[0].add_run("{{ item.description }}")
ap = row[1].paragraphs[0]
ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
ap.add_run("${{ '{:,.2f}'.format(item.amount) }}")
# control row -> {% endfor %}
fee_tbl.add_row().cells[0].paragraphs[0].add_run("{%tr endfor %}")

# total row
trow = fee_tbl.add_row().cells
set_cell_bg(trow[0], LIGHT)
set_cell_bg(trow[1], LIGHT)
tr0 = trow[0].paragraphs[0].add_run("Total Fixed Fee (ex GST)")
tr0.bold = True
tp = trow[1].paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
tp.add_run("${{ '{:,.2f}'.format(fee_total) }}").bold = True

para("", after=4)

# --- optional hourly-rate table (whole block removed when show_hourly is false) ---
tag("{%p if show_hourly %}")
para("Hourly Rates", bold=True, color=ACCENT, after=2)
para("Where additional services are requested beyond the scope above, they will be "
     "charged at the following hourly rates:")

hr_tbl = doc.add_table(rows=1, cols=2)
hr_tbl.style = "Table Grid"
hh = hr_tbl.rows[0].cells
for c, txt in [(hh[0], "Role"), (hh[1], "Rate (AUD/hr, ex GST)")]:
    set_cell_bg(c, "2E75B6")
    rr = c.paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rr.font.size = Pt(10)
hr_tbl.add_row().cells[0].paragraphs[0].add_run("{%tr for r in hourly_items %}")
hrow = hr_tbl.add_row().cells
hrow[0].paragraphs[0].add_run("{{ r.role }}")
hap = hrow[1].paragraphs[0]
hap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hap.add_run("${{ '{:,.2f}'.format(r.rate) }}/hr")
hr_tbl.add_row().cells[0].paragraphs[0].add_run("{%tr endfor %}")
tag("{%p endif %}")

para("", after=2)
gst = para()
gst.add_run("All fees are exclusive of GST. ").italic = True
gst.add_run("{{ payment_terms }}").italic = True

# =====================================================================
# 6. EXCLUSIONS & ASSUMPTIONS  (looped + conditional)
# =====================================================================
tag("{%p if assumptions %}")
para("Assumptions", style="H2")
tag("{%p endif %}")
tag("{%p for a in assumptions %}")
bullet("{{ a }}")
tag("{%p endfor %}")

tag("{%p if exclusions %}")
para("Exclusions", style="H2")
tag("{%p endif %}")
tag("{%p for e in exclusions %}")
bullet("{{ e }}")
tag("{%p endfor %}")

# =====================================================================
# 7. STANDARD TERMS (hyperlinked)
# =====================================================================
para("Standard Terms & Conditions", style="H2")
tp = para()
tp.add_run("This proposal is subject to our standard terms and conditions, available here: ")
add_hyperlink(tp, "https://conceptengineers.com.au/terms", "Standard Terms & Conditions")
tp.add_run(". By accepting this proposal you agree to those terms.")

# =====================================================================
# 8. ACCEPTANCE
# =====================================================================
para("Acceptance", style="H2")
para("We look forward to working with you on this project. To proceed, please sign "
     "below or confirm your acceptance via the secure signing link provided in the "
     "accompanying email.")
para("")
para("Yours faithfully,", after=2)
para("{{ sender_name }}", bold=True, after=0)
para("{{ sender_title }}", after=0)
para("{{ firm_name }}", after=0)
para("{{ sender_email }}  |  {{ sender_phone }}", after=10)

# signature table
sig = doc.add_table(rows=2, cols=2)
sig.style = "Table Grid"
labels = [("Signed (Client)", "Name"), ("Position", "Date")]
fields = [("signature", "name"), ("position", "date")]
for r in range(2):
    for c in range(2):
        cell = sig.rows[r].cells[c]
        lab = cell.paragraphs[0].add_run(labels[r][c] + ":")
        lab.font.size = Pt(9)
        lab.font.color.rgb = GREY
        cell.add_paragraph().add_run(" ")

from pathlib import Path as _Path
out = str(_Path(__file__).parent / "master_proposal_template.docx")
doc.save(out)
print("Saved:", out)
